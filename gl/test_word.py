import copy

import docx
from docx.shared import Pt


def find_second_last_period(string):
    # 找到倒数第二个句号的位置
    second_last_period_index = -1
    last_period_index = -1

    for i, char in enumerate(string):
        if char == '.':
            second_last_period_index = last_period_index
            last_period_index = i

    return second_last_period_index


def find_second_last_period_first_number(string):
    # 找到倒数第二个句号的位置
    second_last_period_index = -1
    last_period_index = -1

    for i, char in enumerate(string):
        if char == '.':
            second_last_period_index = last_period_index
            last_period_index = i

    # 在倒数第二个句号后查找第一个数字的位置
    if second_last_period_index != -1:
        for i in range(second_last_period_index + 1, len(string)):
            if string[i].isdigit():
                return i

    # 如果没有找到数字，则返回-1
    return -1


def find_semicolon_and_first_non_digit(string):
    semicolon_index = string.find(";")

    if semicolon_index != -1:
        # 找到分号后的部分
        after_semicolon = string[semicolon_index + 1:]

        # 移除分号后的空格
        after_semicolon = after_semicolon.lstrip()

        # 找到分号后的第一个非数字字符的位置
        for i, char in enumerate(after_semicolon):
            if not char.isdigit():
                non_digit_index = semicolon_index + 1 + i
                return semicolon_index, non_digit_index

    # 如果没有找到分号或分号后没有非数字字符，则返回-1
    return -1, -1


doc = docx.Document('input.docx')

for paragraph in doc.paragraphs:
    text = copy.deepcopy(paragraph.text)
    paragraph.clear()

    # 删除开头非文本部分
    text = text.lstrip('0123456789. ')

    # 查找关键位置
    period_index = find_second_last_period(text)
    number_index = find_second_last_period_first_number(text)
    semicolon_index, parenthesis_index = find_semicolon_and_first_non_digit(text)

    # 统计句号个数，在删除标号以后小于3个则说明这一行参考文献容易出错，故用蓝色标出
    period_number = text.count('.')

    # 添加其他文本
    text_run = paragraph.add_run(text[: period_index + 1])
    text_run.italic = False
    text_run.bold = False
    if period_number < 3:
        text_run.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0xFF)

    # 设置斜体
    italics_text = text[period_index + 1:number_index]
    italics_run = paragraph.add_run(italics_text)
    italics_run.italic = True
    italics_run.bold = False

    # 添加其他文本
    text_run = paragraph.add_run(text[number_index: semicolon_index + 1])
    text_run.italic = False
    text_run.bold = False

    if semicolon_index != -1 and period_number == 3:
        # 设置粗体
        bold_text = text[semicolon_index + 1:parenthesis_index + 1]
        bold_run = paragraph.add_run(bold_text)
        bold_run.bold = True
        bold_run.italic = False

        # 添加其他文本
        text_run = paragraph.add_run(text[parenthesis_index + 1:])
        text_run.italic = False
        text_run.bold = False

    # 设置段落格式
    paragraph.style = doc.styles['Normal']
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

doc.save('output.docx')
