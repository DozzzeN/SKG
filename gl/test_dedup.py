import copy
from tkinter import messagebox

from docx import Document
from difflib import SequenceMatcher

from docx.shared import RGBColor


def find_text_between_second_and_third_periods(text):
    # 找到第二个句号的位置
    second_period_index = text.find(".", text.find(".") + 1)

    # 找到第三个句号的位置
    third_period_index = text.find(".", second_period_index + 1)

    # 提取第二个句号到第三个句号之间的内容
    if second_period_index != -1 and third_period_index != -1:
        result = text[second_period_index + 1:third_period_index].strip()
        return result
    else:
        return None


# 计算文本相似度
def process_docx(file_path):
    doc = Document(file_path)
    text_set = []

    for i in range(len(doc.paragraphs)):
        # print(i)
        paragraph = copy.deepcopy(doc.paragraphs[i])
        text = paragraph.text.strip()

        # 跳过空白段落
        if not text:
            continue

        title = find_text_between_second_and_third_periods(text)
        if title not in text_set:
            text_set.append(title)
        else:
            doc.paragraphs[i].clear()
            text_run = doc.paragraphs[i].add_run(text)
            text_run.bold = True
            text_run.italic = True
            text_run.font.color.rgb = RGBColor(0, 0, 255)

    doc.save("output.docx")


if __name__ == "__main__":
    input_file = "input.docx"
    process_docx(input_file)
    messagebox.showinfo("提示", "测试结束")